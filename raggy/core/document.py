"""Document processing functionality for text extraction and chunking."""

import hashlib
import re
from pathlib import Path
from typing import Any, Callable, Dict, List, Optional

from ..config.constants import (
    CHUNK_READ_SIZE,
    DEFAULT_CHUNK_OVERLAP,
    DEFAULT_CHUNK_SIZE,
    GLOB_PATTERNS,
    MAX_FILE_SIZE_MB,
)
from ..utils.logging import handle_file_error, log_warning
from ..utils.security import sanitize_error_message, validate_path


class DocumentProcessor:
    """Handles file discovery, text extraction, and chunking operations."""

    def __init__(
        self,
        docs_dir: Path,
        config: Dict[str, Any],
        quiet: bool = False
    ) -> None:
        """Initialize document processor.

        Args:
            docs_dir: Directory containing documents
            config: Configuration dictionary
            quiet: If True, suppress output
        """
        self.docs_dir = docs_dir
        self.config = config
        self.quiet = quiet

        # File type handlers (Strategy pattern)
        self._file_handlers = {
            ".pdf": self._extract_text_from_pdf,
            ".md": self._extract_text_from_md,
            ".docx": self._extract_text_from_docx,
            ".txt": self._extract_text_from_txt,
        }

    def find_documents(self) -> List[Path]:
        """Find all supported documents in docs directory.

        Returns:
            List[Path]: Sorted list of document paths
        """
        if not self.docs_dir.exists():
            if not self.quiet:
                print(f"Creating docs directory: {self.docs_dir}")
            self.docs_dir.mkdir(parents=True, exist_ok=True)
            if not self.quiet:
                print(f"Please add your documentation files to {self.docs_dir}")
            return []

        files = []
        for pattern in GLOB_PATTERNS:
            files.extend(self.docs_dir.glob(pattern))

        return sorted(files)

    def process_document(self, file_path: Path) -> List[Dict[str, Any]]:
        """Process a single document into chunks.

        Args:
            file_path: Path to the document

        Returns:
            List[Dict[str, Any]]: List of document chunks with metadata
        """
        if not self.quiet:
            print(f"Processing: {file_path.relative_to(self.docs_dir)}")

        # Validate file path for security
        if not validate_path(file_path, self.docs_dir):
            log_warning(f"Skipping file outside docs directory: {file_path.name}", quiet=self.quiet)
            return []

        # Check file size limits
        try:
            file_size = file_path.stat().st_size
            if file_size > MAX_FILE_SIZE_MB * 1024 * 1024:
                log_warning(f"Skipping large file (>{MAX_FILE_SIZE_MB}MB): {file_path.name}", quiet=self.quiet)
                return []
        except OSError:
            log_warning(f"Could not check file size for {file_path.name}", quiet=self.quiet)
            return []

        try:
            # Extract text using Strategy pattern
            file_extension = file_path.suffix.lower()
            handler = self._file_handlers.get(file_extension)

            if handler is None:
                if not self.quiet:
                    supported_types = ', '.join(self._file_handlers.keys())
                    print(f"Skipping unsupported file type: {file_path.name}")
                    print(f"Supported types: {supported_types}")
                return []

            text = handler(file_path)

            if not text.strip():
                log_warning(f"No text extracted from {file_path.name}", quiet=self.quiet)
                return []

            # Generate chunks
            chunk_data = self._chunk_text(text)

            # Create document entries
            documents = []
            file_hash = self._get_file_hash(file_path)

            for i, chunk_info in enumerate(chunk_data):
                doc_id = f"{file_path.stem}_{file_hash[:8]}_{i}"

                # Merge chunk metadata with file metadata
                metadata = {
                    "source": str(file_path.relative_to(self.docs_dir)),
                    "chunk_index": i,
                    "total_chunks": len(chunk_data),
                    "file_hash": file_hash,
                    "file_type": file_path.suffix.lower(),
                }
                metadata.update(chunk_info.get("metadata", {}))

                documents.append(
                    {"id": doc_id, "text": chunk_info["text"], "metadata": metadata}
                )

            return documents

        except Exception as e:
            handle_file_error(file_path, "process", e, quiet=self.quiet)
            return []

    def _get_file_hash(self, file_path: Path) -> str:
        """Generate SHA256 hash of file for change detection using streaming for large files.

        Args:
            file_path: Path to the file

        Returns:
            str: SHA256 hash of the file
        """
        hash_sha256 = hashlib.sha256()
        with open(file_path, "rb") as f:
            # Read in chunks to handle large files efficiently
            for chunk in iter(lambda: f.read(CHUNK_READ_SIZE), b""):
                hash_sha256.update(chunk)
        return hash_sha256.hexdigest()

    def _extract_text_template(
        self, file_path: Path, extraction_method: Callable[[Path], str]
    ) -> str:
        """Template method for text extraction with consistent error handling.

        Args:
            file_path: Path to the file
            extraction_method: Method to extract text

        Returns:
            str: Extracted text or empty string on error
        """
        try:
            result = extraction_method(file_path)
            return result.strip() if result else ""
        except ImportError as e:
            # Handle specific import errors (like missing python-docx)
            library = str(e).split("'")[1] if "'" in str(e) else "dependency"
            warning = f"Warning: {library} not available. Cannot read {file_path.name}"
            print(warning)
            return ""
        except Exception as e:
            sanitized_error = sanitize_error_message(str(e))
            print(f"Warning: Could not extract text from {file_path.name}: {sanitized_error}")
            return ""

    def _extract_text_from_pdf(self, file_path: Path) -> str:
        """Extract text from PDF file."""
        return self._extract_text_template(file_path, self._extract_pdf_content)

    def _extract_text_from_md(self, file_path: Path) -> str:
        """Extract text from Markdown file."""
        return self._extract_text_template(file_path, self._extract_md_content)

    def _extract_text_from_docx(self, file_path: Path) -> str:
        """Extract text from Word document (.docx)."""
        return self._extract_text_template(file_path, self._extract_docx_content)

    def _extract_text_from_txt(self, file_path: Path) -> str:
        """Extract text from plain text file."""
        return self._extract_text_template(file_path, self._extract_txt_content)

    def _extract_pdf_content(self, file_path: Path) -> str:
        """Extract content from PDF file."""
        import PyPDF2

        with open(file_path, "rb") as file:
            reader = PyPDF2.PdfReader(file)
            text_parts = []
            for page in reader.pages:
                page_text = page.extract_text()
                if page_text.strip():
                    text_parts.append(page_text)
            return "\n".join(text_parts)

    def _extract_md_content(self, file_path: Path) -> str:
        """Extract content from Markdown file."""
        with open(file_path, "r", encoding="utf-8") as file:
            return file.read()

    def _extract_docx_content(self, file_path: Path) -> str:
        """Extract content from Word document."""
        from docx import Document

        doc = Document(file_path)
        text_parts = []

        # Extract paragraphs
        for paragraph in doc.paragraphs:
            if paragraph.text.strip():
                text_parts.append(paragraph.text.strip())

        # Extract tables
        for table in doc.tables:
            for row in table.rows:
                row_text = []
                for cell in row.cells:
                    if cell.text.strip():
                        row_text.append(cell.text.strip())
                if row_text:
                    text_parts.append(" | ".join(row_text))

        return "\n\n".join(text_parts)

    def _extract_txt_content(self, file_path: Path) -> str:
        """Extract content from plain text file with encoding fallback."""
        # Try UTF-8 first
        try:
            with open(file_path, "r", encoding="utf-8") as file:
                return file.read()
        except UnicodeDecodeError:
            # Fallback to latin-1 for older files
            with open(file_path, "r", encoding="latin-1") as file:
                return file.read()

    def _chunk_text(
        self,
        text: str,
        chunk_size: Optional[int] = None,
        overlap: Optional[int] = None,
        smart: bool = True,
    ) -> List[Dict[str, Any]]:
        """Split text into overlapping chunks with optional smart chunking.

        Args:
            text: Text to chunk
            chunk_size: Size of each chunk
            overlap: Overlap between chunks
            smart: If True, use smart chunking

        Returns:
            List[Dict[str, Any]]: List of chunks with metadata
        """
        chunk_size = chunk_size or self.config["search"].get("chunk_size", DEFAULT_CHUNK_SIZE)
        overlap = overlap or self.config["search"].get("chunk_overlap", DEFAULT_CHUNK_OVERLAP)

        if smart and self.config["chunking"]["smart"]:
            return self._chunk_text_smart(text, chunk_size, overlap)
        else:
            return self._chunk_text_simple(text, chunk_size, overlap)

    def _chunk_text_simple(
        self, text: str, chunk_size: int, overlap: int
    ) -> List[Dict[str, Any]]:
        """Simple chunking for backward compatibility."""
        if len(text) <= chunk_size:
            return [{"text": text, "metadata": {"chunk_type": "simple"}}]

        chunks = []
        start = 0

        while start < len(text):
            end = start + chunk_size

            # Try to break at sentence boundary
            if end < len(text):
                # Look for sentence endings near the chunk boundary
                for i in range(end, max(start + chunk_size - 200, start), -1):
                    if text[i] in ".!?\n":
                        end = i + 1
                        break

            chunk_text = text[start:end].strip()
            if chunk_text:
                chunks.append(
                    {"text": chunk_text, "metadata": {"chunk_type": "simple"}}
                )

            start = end - overlap

        return chunks

    def _chunk_text_smart(
        self, text: str, base_chunk_size: int, overlap: int
    ) -> List[Dict[str, Any]]:
        """Smart chunking with markdown awareness."""
        chunks = []

        # Split by major sections first (headers)
        sections = re.split(r"(^#{1,6}\s+.*$)", text, flags=re.MULTILINE)

        current_header = None
        current_content = ""

        for section in sections:
            if re.match(r"^#{1,6}\s+", section):
                # Process previous section if exists
                if current_content.strip():
                    section_chunks = self._process_section(
                        current_content, current_header, base_chunk_size, overlap
                    )
                    chunks.extend(section_chunks)

                # Start new section
                current_header = section.strip()
                current_content = ""
            else:
                current_content += section

        # Process final section
        if current_content.strip():
            section_chunks = self._process_section(
                current_content, current_header, base_chunk_size, overlap
            )
            chunks.extend(section_chunks)

        # If no headers found, fall back to simple chunking
        if not chunks:
            return self._chunk_text_simple(text, base_chunk_size, overlap)

        return chunks

    def _process_section(
        self, content: str, header: Optional[str], chunk_size: int, overlap: int
    ) -> List[Dict[str, Any]]:
        """Process a section with its header."""
        content = content.strip()
        if not content:
            return []

        target_size = self._determine_target_size(content, chunk_size)
        content_with_header = self._prepend_header_if_needed(content, header)

        if len(content_with_header) <= target_size:
            return [self._create_chunk_metadata(content_with_header, header, 0)]

        return self._split_into_chunks(content_with_header, header, target_size, overlap)

    def _determine_target_size(self, content: str, chunk_size: int) -> int:
        """Determine appropriate chunk size based on content type.

        Args:
            content: Content to analyze
            chunk_size: Base chunk size

        Returns:
            int: Target chunk size
        """
        lines = content.split("\n")
        is_list_content = any(
            line.strip().startswith(("-", "*", "1.")) for line in lines[:5]
        )

        if is_list_content:
            return min(chunk_size, self.config["chunking"]["min_chunk_size"] * 2)

        # Dynamic sizing for regular content
        return min(
            max(len(content) // 3, self.config["chunking"]["min_chunk_size"]),
            self.config["chunking"]["max_chunk_size"],
        )

    def _prepend_header_if_needed(self, content: str, header: Optional[str]) -> str:
        """Prepend header to content if configured to preserve headers.

        Args:
            content: Section content
            header: Section header

        Returns:
            str: Content with header prepended if needed
        """
        if header and self.config["chunking"]["preserve_headers"]:
            return f"{header}\n\n{content}"
        return content

    def _create_chunk_metadata(
        self, text: str, header: Optional[str], chunk_index: int
    ) -> Dict[str, Any]:
        """Create chunk with metadata.

        Args:
            text: Chunk text
            header: Section header
            chunk_index: Index of chunk in section (0 for single chunk)

        Returns:
            Dict[str, Any]: Chunk with metadata
        """
        metadata = {
            "chunk_type": "smart",
            "section_header": header,
            "header_depth": len(re.findall(r"^#", header or "")),
        }

        if chunk_index > 0:
            metadata["section_chunk_index"] = chunk_index

        return {"text": text, "metadata": metadata}

    def _split_into_chunks(
        self, content: str, header: Optional[str], target_size: int, overlap: int
    ) -> List[Dict[str, Any]]:
        """Split content into multiple chunks with overlap.

        Args:
            content: Content to split
            header: Section header
            target_size: Target size for each chunk
            overlap: Number of characters to overlap

        Returns:
            List[Dict[str, Any]]: List of chunks with metadata
        """
        chunks = []
        start = 0
        chunk_index = 0

        while start < len(content):
            end = self._find_chunk_boundary(content, start, target_size)
            chunk_text = content[start:end].strip()

            if chunk_text:
                chunks.append(self._create_chunk_metadata(chunk_text, header, chunk_index))
                chunk_index += 1

            start = end - overlap

        return chunks

    def _find_chunk_boundary(
        self, content: str, start: int, target_size: int
    ) -> int:
        """Find optimal chunk boundary at paragraph or sentence break.

        Args:
            content: Full content
            start: Start position
            target_size: Target chunk size

        Returns:
            int: End position for chunk
        """
        end = start + target_size

        if end >= len(content):
            return len(content)

        # Try paragraph break first
        paragraph_end = self._find_paragraph_break(content, start, end, target_size)
        if paragraph_end > start:
            return paragraph_end

        # Fall back to sentence break
        sentence_end = self._find_sentence_break(content, start, end, target_size)
        if sentence_end > start:
            return sentence_end

        return end

    def _find_paragraph_break(
        self, content: str, start: int, end: int, target_size: int
    ) -> int:
        """Find paragraph break near target end position.

        Args:
            content: Full content
            start: Start position
            end: Target end position
            target_size: Target chunk size

        Returns:
            int: Position of paragraph break or -1 if not found
        """
        # Look back from end for paragraph break
        search_start = max(start + target_size - 300, start)
        for i in range(end, search_start, -1):
            if i > start and content[i - 2 : i] == "\n\n":
                return i
        return -1

    def _find_sentence_break(
        self, content: str, start: int, end: int, target_size: int
    ) -> int:
        """Find sentence break near target end position.

        Args:
            content: Full content
            start: Start position
            end: Target end position
            target_size: Target chunk size

        Returns:
            int: Position of sentence break or -1 if not found
        """
        # Look back from end for sentence boundary
        search_start = max(start + target_size - 200, start)
        for i in range(end, search_start, -1):
            if content[i] in ".!?\n":
                return i + 1
        return -1